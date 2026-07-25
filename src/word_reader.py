
from DocumentReader import DocumentReader
from docx import Document as WordDocument

from document import Document


class WordReader(DocumentReader):
    def __init__(self, file_path: str):
        self.file_path = file_path

    def read(self):

        doc = WordDocument(self.file_path)

        text = "\n".join( p.text  for p in doc.paragraphs  if p.text.strip() )

        return Document(
            text=text,
            source=self.file_path,
            page_count=1
        )